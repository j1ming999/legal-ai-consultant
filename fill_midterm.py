# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json

src = Path(r'C:/Users/16202/Desktop/毕业论文文件/中期检查.docx')
out = src.with_name('中期检查_已填写.docx')
doc = Document(src)
table = doc.tables[0]

TITLE = '基于 RAG 的法律智能问答系统设计与实现'
reading = (
    '已阅读并整理 RAG 检索增强生成、LangChain4J、Spring Boot、Redis 向量检索、MySQL 持久化、JWT 鉴权等相关资料，'
    '重点学习了法律文本按条切分、向量化存储、语义检索与关键词检索融合、流式问答接口设计等内容。'
    '同时收集并导入《中华人民共和国民法典》及相关司法解释、劳动争议、交通事故、消费者权益保护、医疗纠纷等领域的法律法规文本，'
    '为系统知识库建设和问答效果测试提供了资料基础。'
)
progress = (
    '目前毕业设计已完成总体方案设计和主要功能开发。后端采用 Spring Boot 3.5 与 LangChain4J 搭建法律问答服务，'
    '接入通义千问 Qwen-Plus 作为对话模型，使用 text-embedding-v3 对法律文档生成向量并存入 Redis 向量库。'
    '系统已实现法律文档加载、按法条切分、元数据标注、向量检索、关键词检索、加权 RRF 混合检索和查询改写，'
    '能够根据用户问题返回相关法律条文并生成咨询回答。\n'
    '业务功能方面，已完成用户注册、登录、BCrypt 密码加密、JWT 鉴权、会话创建、历史消息保存和删除等功能，'
    '登录用户的对话记录保存到 MySQL，游客模式下对话记录保存在浏览器本地。前端采用 Vue 3 单页面实现，'
    '已具备对话侧边栏、流式输出、Markdown 渲染、法条来源展示、深色模式和移动端适配等功能。'
    '当前系统可以完成从用户提问、检索法条、生成回答到展示引用来源的完整流程，项目主体功能已经基本成型。'
)
problems = (
    '目前主要问题有三点：一是法律问题的口语化表达较多，单纯向量检索有时不能准确命中具体法条；'
    '二是法律法规文本数量增加后，文档切分、去重和检索排序需要进一步优化；'
    '三是系统回答仍需加强约束，避免出现引用不准确或分析过于笼统的情况。\n'
    '针对这些问题，已加入查询改写机制，将口语化问题转换为更适合法律检索的关键词；'
    '检索层采用向量检索与关键词检索结合的方式，并通过 RRF 融合排序提高召回质量；'
    '文档处理时按“第X条”切分，并保存法律名称、章节、条号等元数据，便于法条溯源。'
    '后续将继续扩充测试问题集，比较不同检索策略的效果，完善提示词约束和异常处理，并对前端交互细节进行调整。'
)
teacher = (
    '学生已按计划完成系统需求分析、技术选型、数据库设计和主要功能开发，项目已具备法律问答、法条检索、'
    '用户登录及对话持久化等核心功能。后续应继续加强测试与论文撰写，重点说明 RAG 检索流程、混合检索策略和系统实现效果。'
)

def set_cell(cell, text, font_size=10.5, align=None):
    cell.text = ''
    for i, part in enumerate(text.split('\n')):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.line_spacing = 1.15
        if len(part) > 80:
            p.paragraph_format.first_line_indent = Pt(21)
        run = p.add_run(part)
        run.font.name = '宋体'
        run.font.size = Pt(font_size)

set_cell(table.cell(2, 3), TITLE, 10.5, WD_ALIGN_PARAGRAPH.CENTER)
set_cell(table.cell(3, 1), reading, 10.5)
set_cell(table.cell(4, 1), progress, 10.5)
set_cell(table.cell(5, 1), problems, 10.5)
set_cell(table.cell(6, 1), teacher, 10.5)
doc.save(out)

check_doc = Document(out)
t = check_doc.tables[0]
checks = {
    'exists': out.exists(),
    'size': out.stat().st_size,
    'title': t.cell(2, 3).text,
    'reading_prefix': t.cell(3, 1).text[:80],
    'progress_prefix': t.cell(4, 1).text[:80],
    'problems_prefix': t.cell(5, 1).text[:80],
    'teacher_prefix': t.cell(6, 1).text[:80],
}
print(json.dumps(checks, ensure_ascii=True, indent=2))
